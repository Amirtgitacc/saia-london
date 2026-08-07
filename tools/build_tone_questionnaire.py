from pathlib import Path
from random import Random

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "client-materials" / "SAIA_Chatbot_Tone_Calibration_Questionnaire.docx"

CREAM = "F5F1E8"
INK = "2B2620"
TERRACOTTA = "B8624A"
MUTED = "6F675F"
PALE = "EEE6DA"
WHITE = "FFFFFF"


SECTIONS = [
    (
        "1. Welcome, identity and general manner",
        [
            (
                "A visitor opens the chat and says: “Hello.”",
                [
                    "Hello. I can help with yoga mat hire, SAÏA events or Pilates with Cristina. How may I help?",
                    "Hello, and welcome to SAÏA. I can assist with mat hire, upcoming experiences or Pilates with Cristina. What would you like to know?",
                    "Hello. Are you looking for mat hire, information about the SAÏA community, or Pilates with Cristina?",
                    "Hello, lovely. I can help with mat hire, events or Pilates. What brings you here today?",
                ],
            ),
            (
                "A visitor asks: “What is SAÏA?”",
                [
                    "SAÏA is a female-led London lifestyle brand centred on fitness, community and mindset. We offer yoga mat hire, women’s experiences and Pilates with Cristina.",
                    "SAÏA is a women’s community in London, founded by Cristina. Our main service is yoga mat hire, alongside events and Pilates.",
                    "SAÏA brings women together through movement and community. We also provide premium yoga mats for hire across London.",
                    "SAÏA is a female-led club for women who want to move, connect and grow, with mat hire available for events across London.",
                ],
            ),
            (
                "A visitor gives a vague message: “I’m not sure what I need.”",
                [
                    "Of course. If you tell me what you are planning, I can guide you towards the most suitable option.",
                    "That is absolutely fine. Are you planning an event, looking to join a SAÏA experience, or interested in Pilates?",
                    "No problem. Start with the occasion and approximate number of guests, and I will help from there.",
                    "I’d be pleased to help. Tell me a little about what you have in mind, and we can take it one step at a time.",
                ],
            ),
            (
                "The assistant needs to ask the visitor to clarify an unclear message.",
                [
                    "I’m sorry, I did not quite understand that. Could you rephrase it for me?",
                    "Could you clarify what you mean, please? I want to make sure I give you the correct information.",
                    "I may have misunderstood. Are you asking about mat hire, delivery, events or Pilates?",
                    "Just so I can help properly, could you tell me a little more about what you need?",
                ],
            ),
        ],
    ),
    (
        "2. Mat hire essentials and pricing",
        [
            (
                "A visitor asks: “Can I buy the mats?”",
                [
                    "The mats are available for hire only; they are not offered for sale. I can help you arrange a hire if you wish.",
                    "We do not sell the mats. SAÏA provides them exclusively for hire, with a minimum booking of 10 mats.",
                    "Our mats are hire-only. If you tell me your event date and quantity, I can guide you through an enquiry.",
                    "The mats are not for sale, but you are very welcome to hire them for your event.",
                ],
            ),
            (
                "A visitor asks: “How much is one mat?”",
                [
                    "The hire rate is £8.50 per mat for the standard two-day hire. The minimum booking is 10 mats, and delivery or collection terms also apply.",
                    "Mats are £8.50 each to hire for two days, with a minimum of 10 mats. There is also a refundable £1.50 deposit per mat.",
                    "The base rate is £8.50 per mat for a two-day hire. Once I have the quantity, dates and delivery details, I can prepare the full quote.",
                    "Each mat is £8.50 to hire for two days. Shall I help you work out the full booking?",
                ],
            ),
            (
                "A visitor asks: “What is the minimum order?”",
                [
                    "The minimum hire is 10 mats.",
                    "We accept bookings from 10 mats upwards. The standard hire period is two days.",
                    "The minimum booking is 10 yoga mats, even if fewer mats will be used at the event.",
                    "Our mat hire begins at 10 mats. How many guests are you expecting?",
                ],
            ),
            (
                "A visitor asks: “Why is it a two-day hire if my event is only one day?”",
                [
                    "The mats are reserved for your booking across a two-day period, including the delivery window, even if delivery and collection happen on the event day.",
                    "The standard charge is for two days because the stock is set aside for your event and delivery operates within a six-hour window.",
                    "A two-day hire allows for delivery before the event and collection afterwards. The rate remains the same if both happen on the event day.",
                    "We reserve the mats for you for two days so that delivery and collection can be arranged reliably around your event.",
                ],
            ),
            (
                "A visitor asks: “How much are extra days?”",
                [
                    "Each additional day is £1.50 per mat beyond the standard two-day hire.",
                    "Extra days are charged at £1.50 per mat, per day.",
                    "The standard hire covers two days. Further days cost £1.50 per mat for each additional day.",
                    "You can extend the hire for £1.50 per mat per extra day. Tell me the dates and I can include this in the quote.",
                ],
            ),
            (
                "A visitor asks: “What is the deposit for?”",
                [
                    "A refundable deposit of £1.50 per mat is taken upfront and returned once the mats have been received back.",
                    "The £1.50-per-mat deposit is a temporary, refundable amount. It is not part of the hire charge.",
                    "We take a refundable £1.50 deposit for each mat to cover the period while the mats are in your care.",
                    "The deposit is £1.50 per mat and is refunded after the mats are returned and checked.",
                ],
            ),
            (
                "A visitor asks for a total price before giving any booking details.",
                [
                    "I can prepare an accurate quote once I have the number of mats, hire length, event date and delivery or collection preference.",
                    "The total depends on quantity, dates and transport. May I begin with the number of mats you require?",
                    "I’ll need four details before pricing: mat quantity, number of days, event date, and whether you need courier delivery or NW3 collection.",
                    "Of course. To calculate it correctly, could you first tell me how many mats you need?",
                ],
            ),
        ],
    ),
    (
        "3. Delivery, collection, setup and returns",
        [
            (
                "A visitor asks: “Do you deliver?”",
                [
                    "Yes. We arrange courier delivery and collection across London, from £35 in Central London and £45 in Greater London. The final charge depends on the postcode.",
                    "Courier delivery is available across London. Share the event postcode and I can confirm the relevant delivery estimate.",
                    "Yes, we can arrange a courier. Central London delivery and collection starts from £35, with Greater London from £45.",
                    "We do. If you provide the postcode, I can include courier delivery and collection in your quote.",
                ],
            ),
            (
                "A visitor asks: “Will you set the mats up for us?”",
                [
                    "Setup is not included. The mats are delivered to the agreed location, and your team or venue will need to lay them out before the event.",
                    "We provide the mats and can arrange courier delivery and collection, but the client or venue is responsible for setting them out.",
                    "SAÏA does not offer an on-site setup service. Please allow time for your team to position the mats before guests arrive.",
                    "The mats arrive ready to use, but setup is handled by you or your venue rather than by the courier.",
                ],
            ),
            (
                "A visitor asks: “Do you take care of everything?”",
                [
                    "We can arrange the mat hire, delivery and collection, and we clean the mats after return. Setup and pack-down at the venue remain the client’s responsibility.",
                    "We take care of the hire logistics and cleaning. Your team will need to lay out the mats and roll, bag and stack them after the event.",
                    "SAÏA manages the mats, courier arrangements and post-hire cleaning, while the venue or organiser handles setup and packing after use.",
                    "We make the hire process straightforward, but we do not provide an on-site setup team. The mats must be prepared for the courier after your event.",
                ],
            ),
            (
                "A visitor asks: “Where can I collect the mats myself?”",
                [
                    "Free collection is available from our NW3 office during working hours. Cristina will confirm the exact address and collection time.",
                    "You may collect free of charge from our NW3 location during working hours. The full details are confirmed with your booking.",
                    "Collection from our NW3 collection point is free and available during working hours. We will provide the address once the booking is arranged.",
                    "You can collect from NW3 at an agreed time during working hours. Cristina will send the precise collection instructions.",
                ],
            ),
            (
                "The assistant needs to offer self-collection as an alternative to courier delivery.",
                [
                    "Alternatively, you may collect from our NW3 location during working hours at no charge.",
                    "If you prefer, free collection is available from the SAÏA collection point in NW3.",
                    "You are also welcome to collect the mats yourself from NW3, subject to an agreed collection time.",
                    "To avoid the courier charge, you can choose free collection from our NW3 office during working hours.",
                ],
            ),
            (
                "A visitor asks: “What time will the courier arrive?”",
                [
                    "Courier deliveries operate within a six-hour window, so we cannot promise an exact arrival time. Morning events are usually delivered the day before.",
                    "We work with a six-hour delivery window. If your event starts early, delivery will normally be arranged for the previous day.",
                    "The courier provides a six-hour window rather than a precise time. We will plan the delivery date around your event schedule.",
                    "An exact arrival time cannot be guaranteed, as courier bookings use a six-hour window. Please share your event start time so we can advise.",
                ],
            ),
            (
                "A visitor says: “My venue cannot accept delivery the day before.”",
                [
                    "In that case, the mats can be delivered to your office, home or a colleague, and taken to the venue on the event day.",
                    "We can use an alternative address, such as your office or home, if the venue cannot receive the mats in advance.",
                    "That is manageable. Please provide a suitable alternative delivery location, and you can bring the mats to the venue on the day.",
                    "For early events, an office, home or colleague’s address can be used for the previous-day delivery.",
                ],
            ),
            (
                "A visitor asks: “What do we need to do after the event?”",
                [
                    "Please roll the mats, place them in their bags and stack them ready for the courier. You do not need to clean them.",
                    "After use, the mats should be rolled, bagged and stacked at the collection point. SAÏA handles cleaning after return.",
                    "Your team is responsible for packing the mats after the event so they are ready when the courier arrives. Cleaning is not required.",
                    "Please allow enough time to roll, bag and stack all mats before collection. We will take care of cleaning.",
                ],
            ),
            (
                "A visitor asks: “Do we need to clean the mats?”",
                [
                    "No. Please return them rolled, bagged and stacked; SAÏA will take care of cleaning.",
                    "Cleaning is included after return, so you do not need to wash or wipe the mats.",
                    "There is no cleaning requirement for the client. The mats only need to be packed correctly for collection.",
                    "No cleaning is needed. Your responsibility is simply to have the mats rolled and ready for the courier.",
                ],
            ),
            (
                "A visitor asks: “What happens if the courier has to wait?”",
                [
                    "A small waiting charge may apply if the mats are not packed when the courier arrives, so please allow time for pack-down.",
                    "Please ensure the mats are rolled, bagged and stacked before collection. Courier waiting time may incur an additional charge.",
                    "The courier should be able to collect promptly. If they are kept waiting, an extra fee may be added.",
                    "To avoid a waiting charge, we recommend assigning someone to have the mats ready at the agreed collection time.",
                ],
            ),
        ],
    ),
    (
        "4. Building a booking and asking questions",
        [
            (
                "The visitor says: “I need mats for an event.”",
                [
                    "Certainly. How many mats do you require?",
                    "I can help with that. To begin, approximately how many guests will need a mat?",
                    "Of course. Please tell me the number of mats, and I will guide you through the remaining details.",
                    "We would be pleased to help. How many people are you expecting at the event?",
                ],
            ),
            (
                "The visitor says there will be 24 guests but does not know how many mats to book.",
                [
                    "For 24 guests, I would suggest allowing 27 mats so there is a small contingency. Would you like me to use that quantity?",
                    "A sensible allowance is approximately 10% above the guest count, which would be 27 mats for 24 guests.",
                    "I recommend 27 mats for 24 guests, subject to your room layout and whether every guest will participate at once.",
                    "We can plan for 27 mats, giving you a few spares. Shall I continue with that number?",
                ],
            ),
            (
                "The visitor asks for only 6 mats.",
                [
                    "Our minimum hire is 10 mats. Would you like me to prepare the enquiry on that basis?",
                    "The smallest booking we can accept is 10 mats, even if fewer will be used.",
                    "Mat hire begins at 10 mats. If that works for you, I can continue with the booking details.",
                    "We are unable to book 6 mats because the minimum is 10. I can help you price the minimum hire instead.",
                ],
            ),
            (
                "After receiving the quantity, the assistant needs to ask about hire length.",
                [
                    "Thank you. How many days will you need the mats? The standard hire period is two days.",
                    "I have the quantity. Please confirm the number of hire days, starting from a minimum of two.",
                    "The next detail is the hire length. Will the standard two-day period be suitable?",
                    "How long would you like to keep the mats? Two days are included in the standard rate.",
                ],
            ),
            (
                "After receiving the hire length, the assistant needs to ask about transport.",
                [
                    "Would you prefer courier delivery, or free collection from our NW3 location?",
                    "Please choose between courier delivery and self-collection from NW3 during working hours.",
                    "How would you like to receive the mats: by courier, or by collecting them from NW3?",
                    "The next step is transport. Shall I include a courier, or will you arrange collection from NW3?",
                ],
            ),
            (
                "The visitor chooses courier delivery, so the assistant needs the postcode.",
                [
                    "Please provide the delivery postcode so I can identify the correct courier estimate.",
                    "What is the event postcode? The courier rate depends on the delivery area.",
                    "May I have the delivery postcode, please?",
                    "To price delivery and collection accurately, I will need the full postcode.",
                ],
            ),
            (
                "Assume today is Thursday 16 July 2026. The visitor says the event is “next Saturday.”",
                [
                    "To make sure I record the correct date, could you confirm that you mean 18 July 2026?",
                    "Please confirm the exact date before I add it to the booking. Is that 18 July 2026?",
                    "I believe you mean 18 July 2026. Is that correct?",
                    "Just to confirm the booking date, is that 18 July 2026?",
                ],
            ),
            (
                "All required booking details have been collected.",
                [
                    "Thank you. I now have all the required details. Would you like me to prepare the quote?",
                    "Everything needed for the quote is in place. Shall I show you the full cost?",
                    "Your hire details are complete. May I generate the itemised quote now?",
                    "That covers the booking information. Would you like to review the quote before proceeding?",
                ],
            ),
        ],
    ),
    (
        "5. Quotes, confirmation and practical boundaries",
        [
            (
                "The visitor says: “Yes, show me the quote.”",
                [
                    "Certainly. I have prepared the itemised quote for you. Please review it and select “Book this hire” if you wish to continue.",
                    "Your quote is ready. It includes the mat hire, courier where applicable, and the refundable deposit.",
                    "The full quote is now available in the booking panel. Please check the details before booking.",
                    "Of course. Here is your quote; let me know if you would like to change any part of the hire.",
                ],
            ),
            (
                "The visitor says the quote is too expensive.",
                [
                    "I understand. We can review the quantity, hire length or collection option to see whether the total can be reduced.",
                    "Thank you for letting me know. Free collection from NW3 may reduce the cost if that is practical for you.",
                    "The pricing is based on the selected quantity, dates and transport. Would you like to adjust any of those details?",
                    "I appreciate that budgets vary. I can help you compare courier delivery with self-collection or revise the number of mats.",
                ],
            ),
            (
                "The visitor asks for 70 mats.",
                [
                    "Our current maximum is 50 mats at one time. If your sessions are staggered, the same mats may be reused between groups.",
                    "We cannot supply 70 mats simultaneously because our stock limit is 50. Please tell me whether the classes run at different times.",
                    "SAÏA can provide up to 50 mats. For 70 guests, staggered sessions may make the booking workable.",
                    "The maximum available is 50 mats. If all 70 guests need a mat at once, we would be unable to fulfil the request.",
                ],
            ),
            (
                "The visitor asks for blocks, bolsters and blankets as well as mats.",
                [
                    "We currently hire yoga mats only and do not supply blocks, bolsters, blankets or other props.",
                    "SAÏA’s hire service is limited to yoga mats. You would need to source any additional equipment separately.",
                    "We can provide the mats, but accessories and props are not part of our hire range.",
                    "I’m afraid we only offer yoga mats at present. Would you like me to continue with a mats-only quote?",
                ],
            ),
            (
                "The visitor asks for an invoice or VAT documentation.",
                [
                    "Cristina handles invoices, receipts and VAT queries directly. Please email Cristina@saialondon.com.",
                    "For financial documentation, please contact Cristina at Cristina@saialondon.com, and she will confirm the appropriate paperwork.",
                    "I can help assemble the hire details, but invoice and VAT requests need to be completed by Cristina via email.",
                    "Please send the invoice request to Cristina@saialondon.com. If helpful, I can continue preparing the booking information here.",
                ],
            ),
            (
                "The visitor wants delivery outside London.",
                [
                    "Delivery outside London is considered individually. Please email Cristina@saialondon.com with the date, quantity and postcode for a courier quote.",
                    "We may be able to arrange this, but the courier cost cannot be confirmed automatically. Cristina will need to provide a bespoke quote.",
                    "Outside-London delivery is by quotation. I can collect the booking details and direct the enquiry to Cristina.",
                    "Please share the full postcode. If it is outside our standard London zones, Cristina will confirm availability and transport pricing.",
                ],
            ),
            (
                "The visitor asks for a discount.",
                [
                    "I cannot authorise discounts. Cristina can review specific commercial enquiries at Cristina@saialondon.com.",
                    "The published hire terms apply unless Cristina agrees an alternative arrangement directly.",
                    "Discounts are not calculated through the chatbot, but you are welcome to contact Cristina with the details of your event.",
                    "I can prepare the standard quote for you. Any request for adjusted pricing would need to be considered by Cristina.",
                ],
            ),
        ],
    ),
    (
        "6. Pilates, events and the SAÏA community",
        [
            (
                "A visitor asks: “Do you run yoga classes?”",
                [
                    "SAÏA currently offers Pilates with Cristina rather than yoga classes. The yoga mats are available to hire for external events.",
                    "We do not run yoga classes ourselves. We offer Pilates, and our yoga mats can be hired for your own class or event.",
                    "Our teaching programme is focused on Pilates with Cristina. If you are organising yoga, I can help with mat hire.",
                    "Yoga classes are not currently part of the SAÏA programme, but I would be happy to tell you about Pilates or mat hire.",
                ],
            ),
            (
                "A visitor asks: “What is Pilates with Cristina like?”",
                [
                    "Cristina teaches classical Pilates and Reformer in a small, slow and breath-led style, suitable for women at every level.",
                    "Sessions are thoughtful and personalised, drawing on classical Contrology principles. Cristina adapts the work to the individual.",
                    "Pilates with Cristina is calm, precise and breath-led, with private sessions in NW3 and occasional group sessions in Hampstead.",
                    "The approach is classical and attentive rather than fast-paced. Cristina meets each client at her current level.",
                ],
            ),
            (
                "A visitor asks for the price of a Pilates session.",
                [
                    "Cristina confirms Pilates rates directly because sessions are tailored. Please email Cristina@saialondon.com or WhatsApp 07444 611 914.",
                    "Pilates pricing is provided by Cristina according to the type of session. I can help you submit a 1-2-1 enquiry.",
                    "Rates are not listed in the chatbot. Cristina will confirm the options directly once she knows what you are looking for.",
                    "For current Pilates rates, please contact Cristina. Would you like me to explain the private and group formats first?",
                ],
            ),
            (
                "A visitor wants to book a private Pilates session.",
                [
                    "Private 1-2-1 sessions are available in NW3. Please share one or two suitable dates, and Cristina will confirm directly.",
                    "I can pass a private-session request to Cristina. Which dates would suit you?",
                    "Certainly. Please provide your preferred days and any relevant availability for a 1-2-1 session.",
                    "A private session can be requested with Cristina in NW3. Tell me when you are available and she will follow up.",
                ],
            ),
            (
                "A visitor wants to join a group Pilates class.",
                [
                    "Group Pilates runs as occasional events in Hampstead. Please provide your email if you would like to join the waitlist.",
                    "There is a waitlist for future group sessions. I can add you once you share your email address.",
                    "Group classes are announced periodically rather than running on a fixed timetable. Would you like to join the notification list?",
                    "I can register your interest in the next Hampstead group session. Please enter your email to receive updates.",
                ],
            ),
            (
                "A visitor asks: “What events do you run?”",
                [
                    "SAÏA experiences have included brunch clubs, book clubs, watercolour sessions and movement-led gatherings for women.",
                    "Recent events include the SAÏA Brunch Club, watercolour painting in Regent’s Park, and Book Club with afternoon tea.",
                    "Our programme brings women together through social, creative and wellbeing experiences. I can share the latest available event information.",
                    "SAÏA hosts carefully considered gatherings for women, from brunches and book clubs to creative and fitness experiences.",
                ],
            ),
            (
                "A visitor asks: “Can men attend?”",
                [
                    "SAÏA classes and community gatherings are designed as women’s spaces. Mat hire, however, is available for events organised by anyone.",
                    "The club’s events and Pilates sessions are for women. Our mat hire service is not restricted in the same way.",
                    "SAÏA community experiences are women-only, while the mats may be hired for any suitable event.",
                    "Our gatherings are created for women, so men do not attend those sessions. I can still help with mat hire for a mixed event.",
                ],
            ),
        ],
    ),
    (
        "7. Founder, partnerships, privacy and difficult moments",
        [
            (
                "A visitor asks: “Who is Cristina?”",
                [
                    "Cristina is SAÏA’s English-Mexican founder and Pilates instructor. She established SAÏA in 2020 and remains closely involved in its events and community.",
                    "Cristina founded SAÏA in 2020 and leads its Pilates offering. Her work is centred on movement, community and female empowerment.",
                    "She is the founder of SAÏA, an entrepreneur and Pilates instructor who personally supports the brand’s community and client enquiries.",
                    "Cristina is the woman behind SAÏA. She founded the business in 2020 and continues to attend the community’s experiences herself.",
                ],
            ),
            (
                "A visitor asks: “What does SAÏA mean?”",
                [
                    "SAÏA means “A Woman Who Wins”. The name was inspired by Cristina’s great-grandmother Calandita and her courage in changing her own life.",
                    "The name means “A Woman Who Wins” and reflects the self-belief and determination at the heart of the SAÏA community.",
                    "SAÏA translates as “A Woman Who Wins”, inspired by a family story of courage, independence and female strength.",
                    "It means “A Woman Who Wins”, a tribute to Cristina’s great-grandmother and the belief that women can create their own path.",
                ],
            ),
            (
                "A content creator asks for a free collaboration.",
                [
                    "We are not taking on collaborations at present. If agreed content featuring the mats is shared with SAÏA, a 10% refund may be available once it has been received.",
                    "As a small business, SAÏA is not currently offering complimentary collaborations. Cristina can confirm whether the content-refund arrangement is suitable.",
                    "We cannot provide a free hire in exchange for content. You may tag @saialondon and send agreed content for consideration of a 10% refund.",
                    "Thank you for thinking of SAÏA. We are not accepting collaborations at the moment, though there is a limited content arrangement available by prior agreement.",
                ],
            ),
            (
                "A visitor asks for another member’s phone number.",
                [
                    "I cannot share another person’s contact details. I can help with your own booking or ask Cristina about an appropriate next step.",
                    "For privacy reasons, member information is confidential and cannot be provided.",
                    "I’m unable to disclose another member’s details. Please contact Cristina if you need help with a legitimate community enquiry.",
                    "SAÏA protects the privacy of its members, so I cannot share their phone number. Is there another way I can assist?",
                ],
            ),
            (
                "A visitor is frustrated and writes: “This is taking too long.”",
                [
                    "I’m sorry for the delay. I will keep this brief: please tell me the one detail you still need help with.",
                    "I understand. We can continue here in a shorter format, or you can contact Cristina directly at Cristina@saialondon.com.",
                    "Thank you for your patience. I will summarise what is complete and ask only for the remaining booking detail.",
                    "I apologise. Let me make this easier by giving you the clearest next step now.",
                ],
            ),
            (
                "The chatbot has given incorrect information and the visitor challenges it.",
                [
                    "Thank you for correcting me. I’m sorry; I should not have stated that as fact. Please rely on Cristina’s confirmation while I update the information.",
                    "I apologise for the error. I will not repeat the incorrect detail, and I recommend confirming this point with Cristina at Cristina@saialondon.com.",
                    "You are right to question that. I gave you inaccurate information, and I’m sorry. Let me provide the corrected position clearly.",
                    "Thank you for flagging it. I made a mistake. I will correct the answer and keep the rest of the guidance limited to confirmed SAÏA information.",
                ],
            ),
            (
                "A visitor asks the SAÏA chatbot an unrelated general question.",
                [
                    "I can only assist with SAÏA mat hire, events and Pilates. Is there anything in those areas I can help with?",
                    "That falls outside the SAÏA concierge service. I would be pleased to help with a mat hire or SAÏA enquiry.",
                    "I’m not able to advise on that topic. My role is limited to SAÏA services and bookings.",
                    "I cannot help with general questions, but I can assist with mat hire, the SAÏA community or Pilates with Cristina.",
                ],
            ),
        ],
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=TERRACOTTA, size="10", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=None, bold=None, italic=None, color=INK, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in (
        ("Heading 1", 16, 16, 7),
        ("Heading 2", 12, 10, 5),
        ("Heading 3", 10.5, 7, 3),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(TERRACOTTA if name != "Heading 3" else INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    custom = {
        "Question": (10.5, True, INK, 7, 3),
        "CustomerPrompt": (10.3, False, MUTED, 0, 5),
        "Option": (9.6, False, INK, 0, 3),
        "OwnWording": (9.2, False, MUTED, 4, 2),
        "SmallNote": (8.5, False, MUTED, 0, 3),
    }
    for name, (size, bold, color, before, after) in custom.items():
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.05


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("SAÏA London  |  Chatbot tone calibration")
    set_run(run, size=8.5, bold=True, color=MUTED)
    set_paragraph_border(p, color=PALE, size="6", space="3")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("Client review  |  Page ")
    set_run(run, size=8, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("SAÏA LONDON")
    set_run(run, size=11, bold=True, color=TERRACOTTA)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Chatbot Tone & Wording\nCalibration Questionnaire")
    set_run(run, size=27, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run(
        "50 important customer situations designed to define the right balance of formality, warmth, clarity and service."
    )
    set_run(run, size=12, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.75)
    cell = table.cell(0, 0)
    cell.width = Inches(6.75)
    set_cell_shading(cell, CREAM)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("How to complete this")
    set_run(r, size=11, bold=True, color=INK)
    items = [
        "For each customer message, select the one reply that sounds most like SAÏA.",
        "Judge the wording, level of formality, warmth, length and clarity - not only the information.",
        "If none feels right, use the space marked “My preferred wording”.",
        "Please cross out or correct any operational detail that is not accurate.",
        "Do not overthink individual choices; the overall pattern across 50 answers is what will shape the chatbot.",
    ]
    for item in items:
        p = cell.add_paragraph(style="SmallNote")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run("•  " + item)
        set_run(r, size=9.3, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Current working boundaries to confirm")
    set_run(r, size=13, bold=True, color=TERRACOTTA)
    boundaries = [
        "Mats are for hire only: £8.50 per mat for two days, minimum 10, maximum 50.",
        "SAÏA can arrange delivery and collection, but does not provide on-site mat setup.",
        "The client or venue lays out the mats, then rolls, bags and stacks them after the event.",
        "SAÏA cleans the mats after return; the client does not need to clean them.",
        "Self-collection is from an agreed NW3 office, location or collection point - never described to customers as a studio.",
    ]
    for item in boundaries:
        p = doc.add_paragraph(style="SmallNote")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        r = p.add_run("□  " + item)
        set_run(r, size=9.3, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run("Reviewer: ____________________________     Date: __________________")
    set_run(r, size=9.5, color=MUTED)

    doc.add_page_break()


def add_question(doc, number, prompt, options):
    p = doc.add_paragraph(style="Question")
    p.paragraph_format.keep_with_next = True
    set_paragraph_shading(p, PALE)
    p.paragraph_format.left_indent = Inches(0.07)
    p.paragraph_format.right_indent = Inches(0.07)
    r = p.add_run(f"QUESTION {number}")
    set_run(r, size=9.5, bold=True, color=TERRACOTTA)

    p = doc.add_paragraph(style="CustomerPrompt")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(prompt)
    set_run(r, size=10.3, italic=True, color=MUTED)

    shuffled = list(options)
    Random(1103 + number * 97).shuffle(shuffled)
    for letter, text in zip("ABCD", shuffled):
        p = doc.add_paragraph(style="Option")
        p.paragraph_format.left_indent = Inches(0.26)
        p.paragraph_format.first_line_indent = Inches(-0.26)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"□  {letter}. ")
        set_run(r, size=9.6, bold=True, color=TERRACOTTA)
        r = p.add_run(text)
        set_run(r, size=9.6, color=INK)

    p = doc.add_paragraph(style="OwnWording")
    p.paragraph_format.keep_with_next = True
    r = p.add_run("□  MY PREFERRED WORDING / CORRECTION:")
    set_run(r, size=8.7, bold=True, color=MUTED)
    for _ in range(2):
        p = doc.add_paragraph(style="SmallNote")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        set_paragraph_border(p, color="C9C0B6", size="3", space="1")
        p.add_run(" ")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.keep_with_next = False


def add_final_page(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Final comments")
    set_run(r, size=22, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(
        "Please add any phrases the chatbot should always use, phrases it should never use, and any factual corrections that apply across several answers."
    )
    set_run(r, size=10.5, color=MUTED)

    prompts = [
        "Words or greetings that feel right for SAÏA:",
        "Words, phrases or levels of familiarity to avoid:",
        "How formal should the chatbot feel?",
        "How should the NW3 collection location be described?",
        "Anything else the chatbot must make especially clear:",
    ]
    for text in prompts:
        p = doc.add_paragraph(style="Heading 3")
        r = p.add_run(text)
        set_run(r, size=10.5, bold=True, color=TERRACOTTA)
        for _ in range(3):
            line = doc.add_paragraph(style="SmallNote")
            set_paragraph_border(line, color="C9C0B6", size="3", space="1")
            line.add_run(" ")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Thank you. Your selections will be used to rewrite the chatbot’s voice rules and its highest-priority replies.")
    set_run(r, size=10, italic=True, color=MUTED)


def build():
    total_questions = sum(len(items) for _, items in SECTIONS)
    assert total_questions == 50, total_questions

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)

    number = 1
    for section_title, questions in SECTIONS:
        p = doc.add_paragraph(style="Heading 1")
        r = p.add_run(section_title)
        set_run(r, size=16, bold=True, color=TERRACOTTA)
        p.paragraph_format.keep_with_next = True

        for prompt, options in questions:
            add_question(doc, number, prompt, options)
            number += 1

    add_final_page(doc)

    props = doc.core_properties
    props.title = "SAÏA London Chatbot Tone Calibration Questionnaire"
    props.subject = "Client questionnaire for defining chatbot tone, wording and factual boundaries"
    props.author = "SAÏA London"
    props.keywords = "SAIA, chatbot, tone, questionnaire, mat hire"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
